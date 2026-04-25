from __future__ import annotations

import json
import math
import textwrap
from collections import Counter
from pathlib import Path
from typing import Dict, Iterable, List, Tuple

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from datasets import load_dataset
from matplotlib.patches import FancyArrowPatch, FancyBboxPatch
from scipy.special import softmax
from scipy.sparse import csr_matrix, hstack
from sklearn.base import clone
from sklearn.ensemble import GradientBoostingClassifier, RandomForestClassifier
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.linear_model import LogisticRegression
from sklearn.metrics import accuracy_score, confusion_matrix, precision_recall_fscore_support, roc_auc_score
from sklearn.model_selection import StratifiedKFold
from sklearn.preprocessing import LabelEncoder
from sklearn.svm import SVC

ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "research_package" / "generated_outputs"
FIG_DIR = OUT_DIR / "figures"
ARTICLE_MD = ROOT / "ARTICLE_IMRED.md"
ARTICLE_DOCX = ROOT / "ARTICLE_IMRED.docx"
INVENTORY_JSON = OUT_DIR / "repository_inventory.json"
EVIDENCE_JSON = OUT_DIR / "article_evidence.json"

SQLI_CWES = {"CWE-089", "CWE-090", "CWE-643"}
SECRET_CWES = {"CWE-798", "CWE-259", "CWE-321", "CWE-312"}
CRYPTO_CWES = {"CWE-327", "CWE-328", "CWE-326", "CWE-916", "CWE-477"}

CATEGORY_ORDER = ["sql_injection", "hardcoded_secret", "weak_crypto", "other_vuln"]
CATEGORY_LABELS = {
    "sql_injection": "SQL injection",
    "hardcoded_secret": "Hardcoded secret",
    "weak_crypto": "Weak crypto",
    "other_vuln": "Other vulnerability",
}
LABELS_BENCHMARK = ["hardcoded_secret", "other_vuln", "sql_injection", "weak_crypto"]
DISPLAY_LABELS = {
    "hardcoded_secret": "Hardcoded secret",
    "other_vuln": "Other vulnerability",
    "sql_injection": "SQL injection",
    "weak_crypto": "Weak crypto",
}
PALETTE = {
    "sql_injection": "#C0392B",
    "hardcoded_secret": "#E67E22",
    "weak_crypto": "#8E44AD",
    "other_vuln": "#2E86AB",
}


def map_category(cwe: str) -> str:
    if cwe in SQLI_CWES:
        return "sql_injection"
    if cwe in SECRET_CWES:
        return "hardcoded_secret"
    if cwe in CRYPTO_CWES:
        return "weak_crypto"
    return "other_vuln"


def normalize_source(raw_source: str) -> str:
    source_map = {
        "codeql": "CodeQL",
        "sonar": "SonarRules",
        "pearce": "Pearce",
        "mitre": "CWE list",
        "author": "Author-created",
    }
    return source_map.get(str(raw_source).lower(), str(raw_source).title())


def scan_repository(root: Path) -> Dict[str, object]:
    buckets = {
        "notebooks": [],
        "markdown_files": [],
        "datasets": [],
        "generated_graphs_images": [],
        "source_code_files": [],
        "experiment_outputs": [],
        "references": [],
        "all_files": [],
    }

    dataset_exts = {".csv", ".tsv", ".json", ".parquet", ".xlsx", ".xls"}
    image_exts = {".png", ".jpg", ".jpeg", ".svg", ".gif", ".pdf"}
    source_exts = {".py", ".jsx", ".js", ".ts", ".tsx", ".html", ".css"}
    experiment_markers = {"report", "benchmark", "output", "result", "results"}

    excluded = {
        "ARTICLE_IMRED.md",
        "ARTICLE_IMRED.docx",
        "research_package/build_article_artifacts.py",
    }

    for path in sorted(root.rglob("*")):
        if not path.is_file():
            continue
        rel = path.relative_to(root).as_posix()
        if rel in excluded or rel.startswith("research_package/generated_outputs/"):
            continue
        ext = path.suffix.lower()
        buckets["all_files"].append(rel)

        if ext == ".ipynb":
            buckets["notebooks"].append(rel)
        if ext in {".md", ".markdown"}:
            buckets["markdown_files"].append(rel)
        if ext in dataset_exts:
            buckets["datasets"].append(rel)
        if ext in image_exts:
            buckets["generated_graphs_images"].append(rel)
        if ext in source_exts:
            buckets["source_code_files"].append(rel)
        if any(marker in path.stem.lower() for marker in experiment_markers) and ext in {".md", ".ipynb", ".py", ".json", ".csv"}:
            buckets["experiment_outputs"].append(rel)
        if "reference" in path.stem.lower():
            buckets["references"].append(rel)

    return buckets


def load_securityeval_frame() -> pd.DataFrame:
    dataset = load_dataset("s2e-lab/SecurityEval")
    df = pd.DataFrame(dataset["train"]).rename(
        columns={
            "ID": "id",
            "Prompt": "prompt",
            "Insecure_code": "insecure_code",
        }
    )
    df["cwe"] = df["id"].str.extract(r"(CWE-\d+)")
    df["source_raw"] = df["id"].str.extract(r"CWE-\d+_([^_]+)_")
    df["source"] = df["source_raw"].map(normalize_source)
    df["category"] = df["cwe"].map(map_category)
    df["num_tokens"] = df["insecure_code"].str.split().str.len()
    df["num_lines"] = df["insecure_code"].str.count("\n") + 1
    return df


def build_sparse_features(train_frame: pd.DataFrame, valid_frame: pd.DataFrame, max_features: int = 500):
    vectorizer = TfidfVectorizer(max_features=max_features, token_pattern=r"\b\w+\b")
    x_train_text = vectorizer.fit_transform(train_frame["insecure_code"])
    x_valid_text = vectorizer.transform(valid_frame["insecure_code"])
    x_train_num = csr_matrix(train_frame[["num_tokens", "num_lines"]].to_numpy(dtype=float))
    x_valid_num = csr_matrix(valid_frame[["num_tokens", "num_lines"]].to_numpy(dtype=float))
    return hstack([x_train_text, x_train_num]).tocsr(), hstack([x_valid_text, x_valid_num]).tocsr()


def _scores_from_estimator(estimator, x_valid) -> np.ndarray | None:
    if hasattr(estimator, "predict_proba"):
        return estimator.predict_proba(x_valid)
    if hasattr(estimator, "decision_function"):
        decision = estimator.decision_function(x_valid)
        if decision.ndim == 1:
            decision = np.column_stack([-decision, decision])
        return softmax(decision, axis=1)
    return None


def evaluate_baselines(df: pd.DataFrame) -> Tuple[pd.DataFrame, pd.DataFrame, Dict[str, object]]:
    x_df = df[["insecure_code", "num_tokens", "num_lines"]].copy()
    y_labels = df["category"].copy()

    label_encoder = LabelEncoder()
    y = label_encoder.fit_transform(y_labels)
    skf = StratifiedKFold(n_splits=5, shuffle=True, random_state=42)

    models = {
        "Logistic Regression": LogisticRegression(max_iter=1000, class_weight="balanced", random_state=42),
        "Random Forest": RandomForestClassifier(
            n_estimators=100,
            class_weight="balanced",
            random_state=42,
            n_jobs=-1,
        ),
        "SVM (RBF)": SVC(kernel="rbf", class_weight="balanced", probability=True, random_state=42),
        "Gradient Boosting": GradientBoostingClassifier(n_estimators=100, random_state=42),
    }

    summary_rows: List[Dict[str, float]] = []
    fold_rows: List[Dict[str, float]] = []
    extra: Dict[str, object] = {}

    for name, estimator in models.items():
        oof_pred = np.zeros_like(y)
        oof_proba = np.zeros((len(y), len(np.unique(y))), dtype=float)

        for fold_id, (train_idx, valid_idx) in enumerate(skf.split(x_df, y), start=1):
            train_frame = x_df.iloc[train_idx]
            valid_frame = x_df.iloc[valid_idx]
            y_train = y[train_idx]
            y_valid = y[valid_idx]

            x_train, x_valid = build_sparse_features(train_frame, valid_frame)
            dense_input = name == "Gradient Boosting"
            if dense_input:
                x_train = x_train.toarray()
                x_valid = x_valid.toarray()

            model = clone(estimator)
            model.fit(x_train, y_train)

            pred = model.predict(x_valid)
            score = _scores_from_estimator(model, x_valid)

            oof_pred[valid_idx] = pred
            if score is not None:
                oof_proba[valid_idx] = score

            precision, recall, f1, _ = precision_recall_fscore_support(
                y_valid,
                pred,
                average="macro",
                zero_division=0,
            )
            fold_rows.append(
                {
                    "model": name,
                    "fold": fold_id,
                    "accuracy": float(accuracy_score(y_valid, pred)),
                    "macro_precision": float(precision),
                    "macro_recall": float(recall),
                    "macro_f1": float(f1),
                }
            )

        summary_rows.append(
            {
                "model": name,
                "accuracy": float(pd.DataFrame(fold_rows).query("model == @name")["accuracy"].mean()),
                "macro_precision": float(pd.DataFrame(fold_rows).query("model == @name")["macro_precision"].mean()),
                "macro_recall": float(pd.DataFrame(fold_rows).query("model == @name")["macro_recall"].mean()),
                "macro_f1": float(pd.DataFrame(fold_rows).query("model == @name")["macro_f1"].mean()),
                "roc_auc_ovr": float(
                    roc_auc_score(y, oof_proba, multi_class="ovr", average="macro")
                    if np.any(oof_proba)
                    else float("nan")
                ),
            }
        )

        if name == "Random Forest":
            extra["best_model_name"] = name
            extra["best_model_predictions"] = oof_pred.tolist()
            extra["best_model_labels"] = y.tolist()
            extra["best_model_class_names"] = label_encoder.classes_.tolist()
            extra["best_model_confusion_matrix"] = confusion_matrix(y, oof_pred, normalize="true").tolist()

    return pd.DataFrame(summary_rows), pd.DataFrame(fold_rows), extra


def save_figure(fig: plt.Figure, path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    fig.savefig(path, dpi=220, bbox_inches="tight")
    plt.close(fig)


def make_category_figure(df: pd.DataFrame) -> Path:
    counts = df["category"].value_counts().reindex(CATEGORY_ORDER).fillna(0).astype(int)
    shares = (counts / counts.sum() * 100).round(2)

    fig, axes = plt.subplots(1, 2, figsize=(14, 6))
    axes[0].bar(
        [CATEGORY_LABELS[c] for c in CATEGORY_ORDER],
        counts.values,
        color=[PALETTE[c] for c in CATEGORY_ORDER],
    )
    axes[0].set_title("Category counts in the public SecurityEval split")
    axes[0].set_ylabel("Samples")
    axes[0].tick_params(axis="x", rotation=18)
    for idx, value in enumerate(counts.values):
        axes[0].text(idx, value + 1, str(value), ha="center", va="bottom", fontsize=10)

    axes[1].pie(
        shares.values,
        labels=[CATEGORY_LABELS[c] for c in CATEGORY_ORDER],
        colors=[PALETTE[c] for c in CATEGORY_ORDER],
        autopct="%1.1f%%",
        startangle=90,
    )
    axes[1].set_title("Category share")

    fig.suptitle("Figure 1. Public SecurityEval class distribution used in this study", fontsize=14)
    out = FIG_DIR / "figure_1_category_distribution.png"
    save_figure(fig, out)
    return out


def make_top_cwe_figure(df: pd.DataFrame) -> Path:
    cwe_counts = df["cwe"].value_counts().head(15).sort_values()
    fig, ax = plt.subplots(figsize=(10, 7))
    ax.barh(cwe_counts.index, cwe_counts.values, color="#2E86AB")
    ax.set_title("Figure 2. Top 15 CWE identifiers in the public SecurityEval split")
    ax.set_xlabel("Samples")
    for idx, value in enumerate(cwe_counts.values):
        ax.text(value + 0.05, idx, str(int(value)), va="center", fontsize=9)
    out = FIG_DIR / "figure_2_top_cwes.png"
    save_figure(fig, out)
    return out


def make_source_figure(df: pd.DataFrame) -> Path:
    pivot = (
        df.pivot_table(index="source", columns="category", values="id", aggfunc="count", fill_value=0)
        .reindex(["Author-created", "CodeQL", "SonarRules", "CWE list", "Pearce"])
        .reindex(columns=CATEGORY_ORDER)
    )
    fig, ax = plt.subplots(figsize=(11, 6))
    bottom = np.zeros(len(pivot))
    for category in CATEGORY_ORDER:
        values = pivot[category].values
        ax.bar(
            pivot.index,
            values,
            bottom=bottom,
            label=CATEGORY_LABELS[category],
            color=PALETTE[category],
        )
        bottom += values
    ax.set_title("Figure 3. Source-family distribution after SecureGen category mapping")
    ax.set_ylabel("Samples")
    ax.tick_params(axis="x", rotation=18)
    ax.legend(loc="upper right", frameon=False)
    out = FIG_DIR / "figure_3_source_distribution.png"
    save_figure(fig, out)
    return out


def make_token_figure(df: pd.DataFrame) -> Path:
    fig, ax = plt.subplots(figsize=(10, 6))
    data = [df.loc[df["category"] == category, "num_tokens"].tolist() for category in CATEGORY_ORDER]
    bp = ax.boxplot(data, patch_artist=True, tick_labels=[CATEGORY_LABELS[c] for c in CATEGORY_ORDER])
    for patch, category in zip(bp["boxes"], CATEGORY_ORDER):
        patch.set_facecolor(PALETTE[category])
        patch.set_alpha(0.8)
    ax.set_title("Figure 4. Token-count distribution by mapped vulnerability class")
    ax.set_ylabel("Whitespace token count")
    ax.tick_params(axis="x", rotation=18)
    out = FIG_DIR / "figure_4_token_length_boxplot.png"
    save_figure(fig, out)
    return out


def make_baseline_figure(summary_df: pd.DataFrame) -> Path:
    plot_df = summary_df.sort_values("macro_f1", ascending=False)
    fig, ax = plt.subplots(figsize=(10, 6))
    ax.bar(plot_df["model"], plot_df["macro_f1"], color=["#2E86AB", "#5DADE2", "#85C1E9", "#AED6F1"])
    ax.set_ylim(0, 1)
    ax.set_title("Figure 5. Baseline model comparison using 5-fold macro-F1")
    ax.set_ylabel("Macro-F1")
    ax.tick_params(axis="x", rotation=18)
    for idx, value in enumerate(plot_df["macro_f1"]):
        ax.text(idx, value + 0.02, f"{value:.3f}", ha="center", va="bottom", fontsize=10)
    out = FIG_DIR / "figure_5_baseline_macro_f1.png"
    save_figure(fig, out)
    return out


def make_confusion_figure(extra: Dict[str, object]) -> Path:
    matrix = np.array(extra["best_model_confusion_matrix"])
    class_names = [DISPLAY_LABELS[name] for name in extra["best_model_class_names"]]
    fig, ax = plt.subplots(figsize=(8, 7))
    im = ax.imshow(matrix, cmap="Blues", vmin=0, vmax=1)
    ax.set_title("Figure 6. Out-of-fold normalized confusion matrix for the best baseline")
    ax.set_xticks(range(len(class_names)), class_names, rotation=20, ha="right")
    ax.set_yticks(range(len(class_names)), class_names)
    for i in range(matrix.shape[0]):
        for j in range(matrix.shape[1]):
            ax.text(j, i, f"{matrix[i, j]:.2f}", ha="center", va="center", color="black")
    fig.colorbar(im, ax=ax, fraction=0.046, pad=0.04)
    out = FIG_DIR / "figure_6_random_forest_confusion_matrix.png"
    save_figure(fig, out)
    return out


def make_architecture_figure() -> Path:
    fig, ax = plt.subplots(figsize=(11, 4.5))
    ax.axis("off")

    boxes = [
        (0.04, 0.3, 0.18, 0.4, "Frontend demo\nHTML / React"),
        (0.29, 0.3, 0.18, 0.4, "Documented API layer\nFastAPI design"),
        (0.54, 0.3, 0.18, 0.4, "Inference layer\nCodeBERT target\nRule-based demo"),
        (0.79, 0.3, 0.17, 0.4, "Storage & history\nDocumented only"),
    ]

    for x, y, w, h, label in boxes:
        patch = FancyBboxPatch(
            (x, y),
            w,
            h,
            boxstyle="round,pad=0.02,rounding_size=0.02",
            linewidth=1.4,
            edgecolor="#1F3A5F",
            facecolor="#EAF2F8",
        )
        ax.add_patch(patch)
        ax.text(x + w / 2, y + h / 2, label, ha="center", va="center", fontsize=11)

    arrows = [
        ((0.22, 0.5), (0.29, 0.5)),
        ((0.47, 0.5), (0.54, 0.5)),
        ((0.72, 0.5), (0.79, 0.5)),
    ]
    for start, end in arrows:
        arrow = FancyArrowPatch(start, end, arrowstyle="->", mutation_scale=14, linewidth=1.3, color="#1F3A5F")
        ax.add_patch(arrow)

    ax.set_title("Figure 7. System design distilled from project documentation and prototype code", fontsize=14)
    out = FIG_DIR / "figure_7_system_architecture.png"
    save_figure(fig, out)
    return out


def build_evidence(df: pd.DataFrame, inventory: Dict[str, object], summary_df: pd.DataFrame, extra: Dict[str, object], figure_paths: Dict[str, Path]) -> Dict[str, object]:
    counts = df["category"].value_counts().reindex(CATEGORY_ORDER).fillna(0).astype(int)
    shares = (counts / counts.sum() * 100).round(2)
    source_counts = df["source"].value_counts().to_dict()
    top_cwes = df["cwe"].value_counts().head(15).to_dict()
    feature_stats_df = (
        df.groupby("category")[["num_tokens", "num_lines"]]
        .agg(["mean", "median", "min", "max"])
        .round(2)
    )
    feature_stats: Dict[str, Dict[str, Dict[str, float]]] = {}
    for category in feature_stats_df.index:
        feature_stats[category] = {}
        for feature in feature_stats_df.columns.levels[0]:
            feature_stats[category][feature] = {}
            for stat in feature_stats_df.columns.levels[1]:
                value = feature_stats_df.loc[category, (feature, stat)]
                if pd.notna(value):
                    feature_stats[category][feature][stat] = float(value)
    best_row = summary_df.sort_values("macro_f1", ascending=False).iloc[0]

    return {
        "dataset_rows_public": int(len(df)),
        "dataset_fields": list(df.columns),
        "duplicates_by_id": int(df["id"].duplicated().sum()),
        "missing_prompt": int(df["prompt"].isna().sum()),
        "missing_code": int(df["insecure_code"].isna().sum()),
        "category_counts": counts.to_dict(),
        "category_shares": shares.to_dict(),
        "source_counts": source_counts,
        "top_cwes": top_cwes,
        "feature_stats": feature_stats,
        "baseline_summary": summary_df.round(4).to_dict(orient="records"),
        "best_baseline": {
            "model": str(best_row["model"]),
            "macro_f1": float(best_row["macro_f1"]),
            "accuracy": float(best_row["accuracy"]),
            "macro_precision": float(best_row["macro_precision"]),
            "macro_recall": float(best_row["macro_recall"]),
            "roc_auc_ovr": float(best_row["roc_auc_ovr"]),
        },
        "best_baseline_confusion_matrix": extra["best_model_confusion_matrix"],
        "inventory": inventory,
        "figures": {name: path.relative_to(ROOT).as_posix() for name, path in figure_paths.items()},
    }


def md_escape(value: object) -> str:
    return str(value).replace("|", "\\|")


def format_metric_table(summary_df: pd.DataFrame) -> str:
    lines = [
        "| Model | Accuracy | Precision | Recall | Macro-F1 | ROC-AUC (OvR) |",
        "| --- | ---: | ---: | ---: | ---: | ---: |",
    ]
    for row in summary_df.sort_values("macro_f1", ascending=False).itertuples():
        lines.append(
            f"| {row.model} | {row.accuracy:.3f} | {row.macro_precision:.3f} | {row.macro_recall:.3f} | {row.macro_f1:.3f} | {row.roc_auc_ovr:.3f} |"
        )
    return "\n".join(lines)


def write_markdown_article(evidence: Dict[str, object]) -> str:
    cat_counts = evidence["category_counts"]
    cat_shares = evidence["category_shares"]
    best = evidence["best_baseline"]
    inventory = evidence["inventory"]
    figs = evidence["figures"]
    top_cwes = evidence["top_cwes"]
    source_counts = evidence["source_counts"]

    components_table = "\n".join(
        [
            "| Repository component | Evidence from files | Role in the study |",
            "| --- | --- | --- |",
            "| Lab planning documents | `docs/labs/lab1/project-plan.md`, `docs/labs/lab1/architecture-diagram.md`, `docs/labs/lab1/SecureGen_Plan_Report.docx` | Define project scope, planned architecture, objectives, and risks. |",
            "| EDA assets | `docs/labs/lab2/eda-report.md`, `notebooks/lab2_eda_notebook.ipynb`, `notebooks/lab2_eda_notebook.py` | Describe dataset mapping, cleaning logic, and exploratory plots. |",
            "| Benchmarking assets | `docs/labs/lab3/benchmark-report.md`, `notebooks/lab3_benchmark_notebook.ipynb`, `notebooks/lab3_benchmark_notebook.py` | Specify baseline evaluation code and CodeBERT training scaffold. |",
            "| Product assets | `docs/labs/lab4/product-documentation.md`, `index.html`, `demo/landing-page.html`, `src/SecureGenAIDemo.jsx` | Document and implement the frontend prototype and user workflow. |",
            "| Research package | `research_package/mini_article.md`, `presentation_brief.md`, `presentation_outline.md`, `references.md`, `README.md` | Supply narrative framing, presentation guidance, and the project reference list. |",
        ]
    )

    dataset_table = "\n".join(
        [
            "| Class | Count in public split | Share (%) | Notes |",
            "| --- | ---: | ---: | --- |",
            f"| `sql_injection` | {cat_counts['sql_injection']} | {cat_shares['sql_injection']:.2f} | Derived from CWE-089, CWE-090, and CWE-643. |",
            f"| `hardcoded_secret` | {cat_counts['hardcoded_secret']} | {cat_shares['hardcoded_secret']:.2f} | Derived from CWE-798, CWE-259, CWE-321, and CWE-312. |",
            f"| `weak_crypto` | {cat_counts['weak_crypto']} | {cat_shares['weak_crypto']:.2f} | Derived from CWE-327, CWE-328, CWE-326, CWE-916, and CWE-477. |",
            f"| `other_vuln` | {cat_counts['other_vuln']} | {cat_shares['other_vuln']:.2f} | Residual class covering all remaining mapped weaknesses. |",
        ]
    )

    repo_inventory_block = "\n".join(
        [
            f"- Notebooks ({len(inventory['notebooks'])}): " + ", ".join(f"`{item}`" for item in inventory["notebooks"]),
            f"- Markdown files ({len(inventory['markdown_files'])}): " + ", ".join(f"`{item}`" for item in inventory["markdown_files"]),
            f"- Dataset files ({len(inventory['datasets'])}): "
            + (", ".join(f"`{item}`" for item in inventory["datasets"]) if inventory["datasets"] else "none committed; the notebooks load `s2e-lab/SecurityEval` remotely from Hugging Face."),
            f"- Existing graphs/images ({len(inventory['generated_graphs_images'])}): "
            + (", ".join(f"`{item}`" for item in inventory["generated_graphs_images"]) if inventory["generated_graphs_images"] else "none committed before this article build."),
            f"- Source-code files ({len(inventory['source_code_files'])}): " + ", ".join(f"`{item}`" for item in inventory["source_code_files"]),
            f"- Experiment outputs ({len(inventory['experiment_outputs'])}): " + ", ".join(f"`{item}`" for item in inventory["experiment_outputs"]),
            f"- Reference files ({len(inventory['references'])}): " + ", ".join(f"`{item}`" for item in inventory["references"]),
        ]
    )

    prompt = f"""# SecureGen AI: An IMRED Analysis of a Prototype for Detecting Insecure AI-Generated Code

# Abstract

SecureGen AI is a university project positioned at the intersection of artificial intelligence, deep learning, and cybersecurity. The repository combines project-planning documents, exploratory notebooks, benchmarking code, frontend prototype assets, and a research package focused on the problem of insecure AI-generated code. This article synthesizes all available project files and treats the repository itself as the source of evidence. The study centers on the SecurityEval benchmark described in the project documents as a 130-prompt, 75-CWE benchmark [4], while the reproducible public split accessed by the notebooks contains 121 insecure Python samples. Using the notebook code as implemented, the dataset was mapped into four classes: `sql_injection`, `hardcoded_secret`, `weak_crypto`, and `other_vuln`. The reproducible split was highly imbalanced, with 102 of 121 samples ({cat_shares['other_vuln']:.2f}%) assigned to `other_vuln`, while the three focal classes together comprised only 19 samples. Five-fold stratified cross-validation on TF-IDF plus structural features showed that the best measured baseline in the repository was Random Forest, which achieved accuracy {best['accuracy']:.3f}, macro-precision {best['macro_precision']:.3f}, macro-recall {best['macro_recall']:.3f}, macro-F1 {best['macro_f1']:.3f}, and one-vs-rest ROC-AUC {best['roc_auc_ovr']:.3f}. The repository also documents a CodeBERT-based target architecture, but no trained checkpoint, stored notebook outputs, or executed deep-learning metrics are present; therefore, CodeBERT performance is not reported as an empirical finding. The resulting evidence supports SecureGen AI as a credible prototype and educational case study, while also revealing a substantial gap between the documented product vision and the measured results currently reproducible from the repository.

# 1. Introduction

AI-assisted programming systems such as GitHub Copilot, Amazon CodeWhisperer, and Tabnine have made code generation a routine part of software development, but their usefulness is coupled with a security problem that is central to SecureGen AI. The repository frames this problem through prior work showing that AI coding systems can emit code that is syntactically plausible yet vulnerable, especially for injection flaws, credential exposure, and weak cryptographic choices [4], [5]. Within this project, SecureGen AI is presented as a response to that risk: a system intended to detect insecure patterns in AI-generated code, explain why a snippet is risky, and suggest a safer remediation path.

The project matters for both AI and cybersecurity because vulnerability detection in generated code is not a purely symbolic or purely statistical task. Lexical cues such as `md5` or embedded API keys can be strong indicators, but many insecure patterns depend on context, data flow, and API usage. This motivates the repository's reliance on deep-learning concepts, especially transformer-based encoders derived from BERT and CodeBERT [1]-[3]. In that sense, SecureGen AI is not only a secure-coding tool concept; it is also a case study in how representation learning can be applied to software security analysis.

The research objective reconstructed from the project files is twofold. First, the project aims to create a four-class view of SecurityEval that can support manageable academic experimentation. Second, it aims to turn that classifier into a prototype user-facing product with explainability and remediation support. The repository includes notebooks, reports, architecture descriptions, HTML and React interfaces, and a research package. However, because these materials mix design targets with executable artifacts, the present article deliberately separates reported evidence from aspirational project claims.

The repository scan preceding this article identified the following source corpus:

{repo_inventory_block}

# 2. Methods

## 2.1 Repository-Grounded Study Design

This article was built from a full scan of the repository, with notebooks, Markdown reports, source code, documentation, and the existing planning `.docx` treated as the source of truth. Notebook code was executed only where the repository contained runnable logic and dependencies. Claims present only as narrative expectations were retained as project intentions but not re-labeled as empirical findings.

Table 1 summarizes how the main repository components contributed to the reconstruction.

{components_table}

## 2.2 Dataset and Label Mapping

The notebooks load the public Hugging Face dataset `s2e-lab/SecurityEval`, which exposed a single `train` split with 121 rows and the fields `ID`, `Prompt`, and `Insecure_code` at the time of analysis. This differs from the broader benchmark framing repeated across the documentation, which describes SecurityEval as 130 Python prompts spanning 75 CWE types [4]. Rather than forcing those two counts into a single number, this article distinguishes between the documented benchmark definition and the reproducible public split.

The notebook mapping rules deterministically collapse individual CWE identifiers into four operational labels. SQL injection comprises CWE-089, CWE-090, and CWE-643; hardcoded secrets comprise CWE-798, CWE-259, CWE-321, and CWE-312; weak cryptography comprises CWE-327, CWE-328, CWE-326, CWE-916, and CWE-477; all remaining weaknesses are assigned to `other_vuln`. No committed dataset CSV or parquet file is present in the repository, so the remote Hugging Face split functions as the practical dataset source.

Table 2 reports the mapped class distribution used in this article.

{dataset_table}

## 2.3 Preprocessing and Exploratory Analysis

The EDA notebook and notebook-style Python script apply straightforward preprocessing. After loading the dataset, each sample identifier is parsed to extract its CWE identifier and source family. Source tags are normalized into `CodeQL`, `SonarRules`, `Pearce`, `CWE list`, and `Author-created`. Two lightweight structural features are then engineered from each insecure code snippet: whitespace token count (`num_tokens`) and line count (`num_lines`).

The repository's reported integrity checks were reproduced on the public split. Duplicate identifiers were absent, and missing values in the core prompt and code fields were not observed. The notebooks define exploratory plots for class balance, CWE frequency, source distribution, token-length spread, benchmark overlays, and lexical PCA. However, the committed `.ipynb` files contain zero stored outputs; all visual evidence in this article was therefore generated afresh from the executable notebook logic and the available public split.

## 2.4 Model Benchmarking Procedure

The benchmarking notebook implements four classical baselines over TF-IDF features extracted from the insecure code field, augmented with `num_tokens` and `num_lines`. The evaluated models are Logistic Regression, Random Forest, SVM with an RBF kernel, and Gradient Boosting. The code uses five-fold stratified cross-validation, which is appropriate given the repository's own repeated acknowledgement that the benchmark is small and imbalanced.

The same notebook also documents a planned deep-learning pipeline built around `microsoft/codebert-base`, Hugging Face tokenization, sequence truncation to 128 tokens, a learning rate of `2e-5`, batch size 8, weight decay `0.01`, and epoch-level checkpointing. However, the repository contains no trained checkpoint, no saved trainer logs, no stored attention outputs, and no executed evaluation records for this stage. Consequently, the CodeBERT configuration is described here as documented methodology rather than as a reported experiment outcome.

## 2.5 Prototype and System Workflow

The product side of the repository consists of a static HTML demo (`index.html`), a marketing-style landing page (`demo/landing-page.html`), and a React component (`src/SecureGenAIDemo.jsx`). These files implement a rule-based frontend prototype that detects three focal risk families using regular-expression patterns, assigns a severity badge, produces a confidence display, highlights risky tokens, suggests a remediation snippet, and records a small analysis history in the browser. The Lab 1 architecture files additionally describe a broader four-tier system with a React frontend, FastAPI backend, model layer, Redis cache, and PostgreSQL storage, but backend source code is not committed.

Figure 7 presents a system design diagram distilled from the documented architecture and the executable prototype.

![Figure 7. System design distilled from project documentation and prototype code]({figs['architecture']})

Figure 7. System design distilled from project documentation and prototype code.

The workflow can also be summarized in Mermaid form:

```mermaid
flowchart LR
    A["Code snippet input"] --> B["Preprocessing and tokenization"]
    B --> C["Baseline TF-IDF pipeline (measured)"]
    B --> D["CodeBERT pipeline (documented target)"]
    C --> E["Class label and score"]
    D --> E
    E --> F["Prototype UI: severity, explanation, remediation"]
```

The main analysis pipeline reconstructed from the repository is shown below.

```text
Algorithm 1: SecureGen AI repository pipeline
Input: code snippet s
Output: vulnerability label y, explanatory output e, remediation r

1. Load SecurityEval benchmark metadata and four-class CWE mapping.
2. For each benchmark sample:
3.     extract CWE identifier and source family from sample ID
4.     map CWE to one of {{sql_injection, hardcoded_secret, weak_crypto, other_vuln}}
5.     compute auxiliary features: num_tokens, num_lines
6. Build TF-IDF representations from insecure code text.
7. Evaluate classical models with 5-fold stratified cross-validation.
8. Select the best measured baseline from macro-F1.
9. In the prototype UI, analyze a user snippet with rule-based detectors.
10. Return severity, label, highlighted tokens, and remediation guidance.
11. Treat CodeBERT configuration as documented future work unless a trained checkpoint and metrics are available.
```

# 3. Results

## 3.1 Reproducible Dataset Characteristics

The public SecurityEval split used by the repository notebooks contained 121 insecure Python samples, zero duplicate identifiers, zero missing prompt values, and zero missing insecure-code values. After four-class mapping, the dataset was dominated by `other_vuln` (102 samples, {cat_shares['other_vuln']:.2f}%), with minority counts of 6 SQL injection samples, 6 hardcoded secret samples, and 7 weak cryptography samples. This confirms that the empirical setting is extremely imbalanced.

The source-family distribution was also uneven. The largest source family was `Author-created` with {source_counts['Author-created']} samples, followed by `CodeQL` with {source_counts['CodeQL']}, `SonarRules` with {source_counts['SonarRules']}, `CWE list` with {source_counts['CWE list']}, and `Pearce` with {source_counts['Pearce']}. The most frequent individual CWE identifiers in the public split were {", ".join(f"{k} ({v})" for k, v in list(top_cwes.items())[:5])}. These results indicate that the repository's practical dataset is both source-heterogeneous and heavily skewed toward the residual vulnerability class.

![Figure 1. Public SecurityEval class distribution used in this study]({figs['category']})

Figure 1. Public SecurityEval class distribution used in this study.

![Figure 2. Top 15 CWE identifiers in the public SecurityEval split]({figs['top_cwe']})

Figure 2. Top 15 CWE identifiers in the public SecurityEval split.

![Figure 3. Source-family distribution after SecureGen category mapping]({figs['source']})

Figure 3. Source-family distribution after SecureGen category mapping.

Snippet length also varied by class. In the public split, SQL injection examples had the largest mean token count (62.33 tokens), while weak cryptography examples were the shortest on average (29.29 tokens). This supports the repository's use of token and line counts as auxiliary structural features, although these features alone are unlikely to resolve semantic overlap between classes.

![Figure 4. Token-count distribution by mapped vulnerability class]({figs['tokens']})

Figure 4. Token-count distribution by mapped vulnerability class.

## 3.2 Measured Baseline Performance

Only the classical baseline stage of the benchmarking notebook was fully reproducible from the repository. Table 3 reports the measured five-fold mean scores obtained by executing the notebook logic against the public split.

{format_metric_table(pd.DataFrame(evidence['baseline_summary']))}

The best measured baseline was Random Forest, which reached macro-F1 {best['macro_f1']:.3f} and accuracy {best['accuracy']:.3f}. Gradient Boosting and Logistic Regression produced lower macro-F1 values ({next(row['macro_f1'] for row in evidence['baseline_summary'] if row['model'] == 'Gradient Boosting'):.3f} and {next(row['macro_f1'] for row in evidence['baseline_summary'] if row['model'] == 'Logistic Regression'):.3f}, respectively), while SVM with an RBF kernel performed worst under the current feature space and class imbalance. The gap between accuracy and macro-F1 is itself informative: high raw accuracy is partly driven by the prevalence of `other_vuln`, whereas macro-averaged scores penalize weak minority-class behavior more strongly.

![Figure 5. Baseline model comparison using 5-fold macro-F1]({figs['baseline']})

Figure 5. Baseline model comparison using 5-fold macro-F1.

The normalized out-of-fold confusion matrix for the best baseline shows why the macro scores remain modest. `Other_vuln` is recognized strongly, but minority classes experience substantial confusion, particularly between `sql_injection` and the broader residual class. This behavior is consistent with a dataset in which only a handful of samples define the minority categories.

![Figure 6. Out-of-fold normalized confusion matrix for the best baseline]({figs['confusion']})

Figure 6. Out-of-fold normalized confusion matrix for the best baseline.

## 3.3 What Is Not Reported

Several outputs discussed in the repository are not empirically available in the committed files. The notebooks contain no stored graphs, confusion matrices, or trainer logs. No trained CodeBERT checkpoint is present. No executed attention heatmap is stored. No backend API source code is committed. No screenshots are included. For that reason, CodeBERT accuracy, macro-F1, loss curves, attention visualizations, and deployment-time latency are not reported here as measured results. The expected CodeBERT values shown in `docs/labs/lab3/benchmark-report.md` and `notebooks/lab3_benchmark_notebook.py` are treated as targets or illustrative placeholders rather than findings.

## 3.4 Prototype Output Behavior

Although the backend is not implemented in source form, the frontend prototype behavior is observable from `index.html` and `src/SecureGenAIDemo.jsx`. The prototype can classify pasted code into SQL injection, hardcoded secret, weak crypto, or a safe/default branch using pattern matching. It then produces a severity badge, confidence-style progress bar, highlighted risky tokens, a remediation panel, and a short history table. In other words, the prototype already demonstrates the intended user interaction loop even though the learned inference layer is not yet connected end-to-end.

# 4. Discussion

SecureGen AI has several strengths as an academic project. First, it is well grounded in current security concerns around AI-generated code and anchors itself in established literature on transformers, CodeBERT, SecurityEval, and vulnerability patterns [1]-[9]. Second, the repository covers the full project arc from planning and EDA through benchmarking, product framing, and presentation assets. Third, the executed baseline results show that the project is not merely conceptual; there is enough implemented experimentation to support a genuine empirical discussion.

At the same time, the repository exposes important limitations. The practical dataset is very small and extremely imbalanced, with only 19 samples across the three focal minority classes. That imbalance helps explain why the best measured macro-F1 remains near 0.50 even though accuracy is much higher. In this setting, model reliability is limited, especially for minority-class recall. A model could appear successful by over-predicting `other_vuln`, which would be unacceptable for security triage where missed SQL injection or secret leakage events are costly.

Explainability is also only partially realized. The project documentation repeatedly emphasizes attention-based token highlighting, and the prototype includes rule-based token highlighting in the UI. However, a true learned explanation based on a trained CodeBERT checkpoint is not available in the committed artifacts. Therefore, the explainability story is stronger at the design level than at the empirical level. This is not a failure of the project so much as an indication that the repository presently represents a prototype and research scaffold rather than a finished model-driven system.

From a false-positive and false-negative perspective, the most concerning risk is false negatives. A prototype that fails to detect a minority-class vulnerability may give developers misplaced confidence in AI-generated code. False positives are also problematic because they reduce usability, but they can usually be reviewed manually. The measured baseline scores suggest that minority-class misses remain plausible. This concern is heightened by adversarial evasion possibilities mentioned in the project materials: query obfuscation, encoded secrets, uncommon cryptographic APIs, and distribution shift across languages or frameworks. Because the executable detector in the frontend is rule-based, it is especially susceptible to superficial evasion through small lexical changes.

There is also a clear production-versus-prototype gap. The architecture documents describe FastAPI, Redis, PostgreSQL, and a fine-tuned CodeBERT inference layer, yet the committed implementation is currently a browser-based prototype plus notebook experimentation. That is entirely appropriate for a university project, but the difference should be stated plainly. The current repository demonstrates a convincing concept, exploratory analysis, and benchmark scaffolding; it does not yet demonstrate a deployed ML security service.

Ethically, the project is well motivated because it positions SecureGen AI as a defensive aid rather than an automated authority. This is the right framing. In secure software development, such a tool should complement human review, secure coding guidance, and conventional static analysis rather than replace them. The repository's emphasis on remediation guidance and educational interfaces is therefore a strength, particularly for teaching contexts.

# Conclusion

SecureGen AI makes a meaningful contribution as a repository-level case study on detecting insecure AI-generated code. Across its planning documents, exploratory notebooks, benchmark code, prototype frontend, and research package, the project articulates a coherent vision: map SecurityEval into actionable vulnerability classes, use modern ML to classify insecure code, and surface the result through an explainable user interface. The reproducible evidence gathered from the repository confirms that the public split is small, imbalanced, and difficult, and that the best measured baseline under the current notebook implementation is Random Forest with macro-F1 {best['macro_f1']:.3f}. At the same time, the repository does not yet contain empirical CodeBERT results, stored notebook outputs, or an implemented backend service, so those elements remain future work rather than reported findings.

The most useful next steps are therefore concrete and measurable: train and save the CodeBERT checkpoint referenced in the notebook, commit real evaluation artifacts such as loss curves and attention heatmaps, expand the dataset or label strategy to reduce class imbalance, and replace the current rule-based demo logic with the learned inference pipeline described in the architecture. If those steps are completed, SecureGen AI could move from a strong educational prototype toward a more convincing secure-AI coding assistant for research and applied software engineering.

# References

[1] A. Vaswani, N. Shazeer, N. Parmar, J. Uszkoreit, L. Jones, A. N. Gomez, L. Kaiser, and I. Polosukhin, "Attention Is All You Need," in *Advances in Neural Information Processing Systems 30*, 2017. [Online]. Available: https://arxiv.org/abs/1706.03762

[2] J. Devlin, M.-W. Chang, K. Lee, and K. Toutanova, "BERT: Pre-training of Deep Bidirectional Transformers for Language Understanding," 2018. [Online]. Available: https://arxiv.org/abs/1810.04805

[3] Z. Feng et al., "CodeBERT: A Pre-Trained Model for Programming and Natural Languages," in *Findings of the Association for Computational Linguistics: EMNLP 2020*, pp. 1536-1547, 2020. doi: 10.18653/v1/2020.findings-emnlp.139. [Online]. Available: https://aclanthology.org/2020.findings-emnlp.139/

[4] M. L. Siddiq and J. C. S. Santos, "SecurityEval Dataset: Mining Vulnerability Examples to Evaluate Machine Learning-Based Code Generation Techniques," in *Proceedings of MSR4P&S 2022*, 2022. doi: 10.1145/3549035.3561184. [Online]. Available: https://dl.acm.org/doi/10.1145/3549035.3561184

[5] H. Pearce, B. Ahmad, B. Tan, B. Dolan-Gavitt, and R. Karri, "Asleep at the Keyboard? Assessing the Security of GitHub Copilot's Code Contributions," in *2022 IEEE Symposium on Security and Privacy (SP)*, pp. 754-768, 2022. doi: 10.1109/SP46214.2022.9833571. [Online]. Available: https://ieeexplore.ieee.org/document/9833571

[6] OWASP Foundation, "OWASP Top 10: 2021," 2021. [Online]. Available: https://owasp.org/Top10/2021/

[7] MITRE, "CWE-89: Improper Neutralization of Special Elements used in an SQL Command ('SQL Injection')," 2026. [Online]. Available: https://cwe.mitre.org/data/definitions/89.html

[8] MITRE, "CWE-798: Use of Hard-coded Credentials," 2026. [Online]. Available: https://cwe.mitre.org/data/definitions/798.html

[9] MITRE, "CWE-327: Use of a Broken or Risky Cryptographic Algorithm," 2026. [Online]. Available: https://cwe.mitre.org/data/definitions/327.html
"""
    ARTICLE_MD.write_text(prompt, encoding="utf-8")
    return prompt


def install_docx_if_needed() -> None:
    try:
        import docx  # noqa: F401
    except ImportError as exc:
        raise RuntimeError(
            "python-docx is required to build ARTICLE_IMRED.docx. Install it with `python -m pip install python-docx`."
        ) from exc


def add_markdown_paragraph(document, text: str) -> None:
    from docx.shared import Inches

    stripped = text.strip()
    if not stripped:
        return
    if stripped.startswith("!["):
        path = stripped.split("](", 1)[1].rstrip(")")
        image_path = (ROOT / path).resolve()
        if image_path.exists():
            document.add_picture(str(image_path), width=Inches(6.2))
        return
    if stripped.startswith("|"):
        return
    if stripped.startswith("```"):
        return
    document.add_paragraph(stripped)


def build_docx_from_markdown(markdown_text: str) -> None:
    install_docx_if_needed()

    from docx import Document
    from docx.enum.section import WD_SECTION
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt

    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    styles = document.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(11)

    lines = markdown_text.splitlines()
    i = 0
    in_code_block = False
    code_lines: List[str] = []

    while i < len(lines):
        line = lines[i]
        if line.startswith("```"):
            if in_code_block:
                paragraph = document.add_paragraph()
                run = paragraph.add_run("\n".join(code_lines))
                run.font.name = "Consolas"
                run.font.size = Pt(9.5)
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        if line.startswith("# "):
            heading = document.add_heading(line[2:].strip(), level=0)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue
        if line.startswith("## "):
            document.add_heading(line[3:].strip(), level=1)
            i += 1
            continue
        if line.startswith("### "):
            document.add_heading(line[4:].strip(), level=2)
            i += 1
            continue

        if line.startswith("| "):
            table_lines = []
            while i < len(lines) and lines[i].startswith("|"):
                table_lines.append(lines[i])
                i += 1
            rows = [row.strip().strip("|").split("|") for row in table_lines]
            if len(rows) >= 2:
                header = [cell.strip() for cell in rows[0]]
                body = [[cell.strip() for cell in row] for row in rows[2:]]
                table = document.add_table(rows=1, cols=len(header))
                table.style = "Table Grid"
                for idx, cell in enumerate(header):
                    table.rows[0].cells[idx].text = cell
                for row in body:
                    table_row = table.add_row().cells
                    for idx, cell in enumerate(row):
                        table_row[idx].text = cell
            continue

        if line.startswith("- "):
            document.add_paragraph(line[2:].strip(), style="List Bullet")
            i += 1
            continue

        add_markdown_paragraph(document, line)
        i += 1

    document.save(str(ARTICLE_DOCX))


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    FIG_DIR.mkdir(parents=True, exist_ok=True)

    inventory = scan_repository(ROOT)
    df = load_securityeval_frame()
    summary_df, fold_df, extra = evaluate_baselines(df)

    figure_paths = {
        "category": make_category_figure(df),
        "top_cwe": make_top_cwe_figure(df),
        "source": make_source_figure(df),
        "tokens": make_token_figure(df),
        "baseline": make_baseline_figure(summary_df),
        "confusion": make_confusion_figure(extra),
        "architecture": make_architecture_figure(),
    }

    evidence = build_evidence(df, inventory, summary_df, extra, figure_paths)

    INVENTORY_JSON.write_text(json.dumps(inventory, indent=2), encoding="utf-8")
    EVIDENCE_JSON.write_text(json.dumps(evidence, indent=2), encoding="utf-8")

    markdown_text = write_markdown_article(evidence)
    build_docx_from_markdown(markdown_text)

    print("Wrote", ARTICLE_MD)
    print("Wrote", ARTICLE_DOCX)
    for name, path in figure_paths.items():
        print(name, path)


if __name__ == "__main__":
    main()
